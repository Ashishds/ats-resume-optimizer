from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import re

def txt_to_docx_bytes(text: str) -> bytes:
    doc = Document()
    for line in text.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()

def txt_to_pdf_bytes(text: str) -> bytes:
    """Convert text to PDF bytes using ReportLab with professional resume formatting"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, 
                          topMargin=0.75*inch, 
                          bottomMargin=0.75*inch,
                          leftMargin=0.75*inch,
                          rightMargin=0.75*inch)
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Create professional resume styles
    name_style = ParagraphStyle(
        'NameStyle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=6,
        textColor=colors.darkblue,
        fontName='Helvetica-Bold',
        alignment=1  # Center alignment
    )
    
    contact_style = ParagraphStyle(
        'ContactStyle',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=12,
        textColor=colors.black,
        fontName='Helvetica',
        alignment=1  # Center alignment
    )
    
    section_heading_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=8,
        spaceBefore=12,
        textColor=colors.darkblue,
        fontName='Helvetica-Bold',
        borderWidth=1,
        borderColor=colors.darkblue,
        borderPadding=4,
        backColor=colors.lightgrey
    )
    
    job_title_style = ParagraphStyle(
        'JobTitle',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=2,
        textColor=colors.black,
        fontName='Helvetica-Bold'
    )
    
    company_style = ParagraphStyle(
        'Company',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=2,
        textColor=colors.black,
        fontName='Helvetica'
    )
    
    date_style = ParagraphStyle(
        'Date',
        parent=styles['Normal'],
        fontSize=9,
        spaceAfter=4,
        textColor=colors.grey,
        fontName='Helvetica',
        alignment=2  # Right alignment
    )
    
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=3,
        leftIndent=15,
        textColor=colors.black,
        fontName='Helvetica'
    )
    
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=4,
        textColor=colors.black,
        fontName='Helvetica'
    )
    
    # Build content
    story = []
    lines = text.splitlines()
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            story.append(Spacer(1, 3))
            i += 1
            continue
        
        # Check for name (usually first line, all caps, short)
        if i == 0 and line.isupper() and len(line) < 50 and len(line.split()) <= 4:
            story.append(Paragraph(line, name_style))
            story.append(Spacer(1, 6))
        
        # Check for contact info (email, phone, location)
        elif '@' in line or '+' in line or any(word in line.lower() for word in ['email', 'phone', 'address', 'location']):
            story.append(Paragraph(line, contact_style))
        
        # Check for section headings
        elif any(section in line.upper() for section in ['EXPERIENCE', 'EDUCATION', 'SKILLS', 'PROJECTS', 'SUMMARY', 'OBJECTIVE', 'ACHIEVEMENTS', 'CERTIFICATIONS', 'LANGUAGES']):
            story.append(Paragraph(line.upper(), section_heading_style))
            story.append(Spacer(1, 6))
        
        # Check for bullet points
        elif line.startswith(('•', '-', '*', '◦', '·')):
            clean_line = re.sub(r'^[•\-*◦·]\s*', '', line)
            story.append(Paragraph(f"• {clean_line}", bullet_style))
        
        # Check for job titles (usually followed by company and dates)
        elif len(line.split()) <= 4 and not any(char in line for char in ['@', '+', '|', '•', '-', '*']):
            # Check if next line might be company
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if any(indicator in next_line for indicator in ['|', 'at', 'Company', 'Inc', 'LLC', 'Corp']):
                    story.append(Paragraph(line, job_title_style))
                else:
                    story.append(Paragraph(line, normal_style))
            else:
                story.append(Paragraph(line, normal_style))
        
        # Check for company names (contains | or common company indicators)
        elif '|' in line or any(indicator in line for indicator in ['at', 'Company', 'Inc', 'LLC', 'Corp', 'Ltd']):
            story.append(Paragraph(line, company_style))
        
        # Check for dates (contains year or date patterns)
        elif re.search(r'\b(19|20)\d{2}\b', line) or any(month in line.lower() for month in ['jan', 'feb', 'mar', 'apr', 'may', 'jun', 'jul', 'aug', 'sep', 'oct', 'nov', 'dec']):
            story.append(Paragraph(line, date_style))
        
        # Regular paragraph
        else:
            story.append(Paragraph(line, normal_style))
        
        i += 1
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()
